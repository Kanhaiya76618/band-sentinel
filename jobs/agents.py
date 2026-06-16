"""
Aegis — the five job agents.

Same shape as backend/agents/roster.py: each agent turns real computation
(provider search, fit scoring, resume tailoring, application packaging) into a
typed RoomMessage and posts it to the bus. The LLM supplies only the room's
natural-language line; the logic underneath is real Python.

    @observer   parse criteria/resume -> SearchProfile
    @validator  Adzuna search + fit scoring -> JobMatches      (the skeptic/ranker)
    @commander  present ranked matches, gate on a human
    @tailor     rewrite the resume to a chosen posting (md + PDF + DOCX)
    @applier    submit (real email-apply) or queue (apply link) — honestly
"""
from __future__ import annotations

import os
import re
from pathlib import Path


def _offline() -> bool:
    """OfflineLLM echoes its prompt, so skip generative calls when offline."""
    return os.getenv("LLM_MODE", "offline").lower() == "offline"

from backend.contracts import Intent, RoomMessage
from backend.llm import LLMClient
from backend.store import DATA_DIR

from .contracts import (
    Application, JobDecision, JobMatch, JobMatches, SearchProfile, TailorResult,
)
from .providers import JobSearchProvider
from .resume import SKILL_VOCAB

ARTIFACTS = DATA_DIR / "artifacts"


class JobAgent:
    def __init__(self, agent_id: str, llm: LLMClient):
        self.id = agent_id
        self.llm = llm

    def _voice(self, tag: str, fallback: str) -> str:
        return self.llm.complete(
            system=f"You are {self.id}, a job-search agent. One terse sentence.",
            user=fallback, role=self.id, tag=tag,
        ) or fallback


# --------------------------------------------------------------------------- #
class Observer(JobAgent):
    framework, provider = "LangGraph", "aiml"

    def build_profile(self, *, entry_mode: str, query: str,
                      location: str | None = None, resume: dict | None = None) -> RoomMessage:
        if entry_mode == "resume" and resume:
            prof = SearchProfile(
                entry_mode="resume", query=query or (resume.get("titles") or ["your profile"])[0],
                titles=resume.get("titles", []), skills=resume.get("skills", []),
                seniority=resume.get("seniority"), location=location,
                summary=resume.get("summary"), resume_text=resume.get("resume_text"),
                source="resume",
            )
            line = (f"Parsed your resume: {len(prof.skills)} skills, "
                    f"top title '{(prof.titles or ['n/a'])[0]}'. Searching matching roles. "
                    "@validator, find current openings.")
        elif entry_mode == "company":
            prof = SearchProfile(entry_mode="company", query=query, company=query,
                                 location=location, titles=[], skills=[], source="criteria")
            line = f"Target company: {query}. @validator, pull their current openings."
        else:  # field
            prof = SearchProfile(entry_mode="field", query=query, titles=[query],
                                 location=location, skills=[], source="criteria")
            line = f"Field search: {query}. @validator, find current matching roles."
        return RoomMessage.of(self.id, Intent.SEARCH_PROFILE, self._voice("profile", line),
                              mentions=["@validator"], payload_model=prof)


# --------------------------------------------------------------------------- #
class Validator(JobAgent):
    """Finds CURRENT real postings and scores fit vs the profile (the skeptic)."""
    framework, provider = "CrewAI", "featherless"

    def search_and_rank(self, prof: SearchProfile, provider: JobSearchProvider,
                        limit: int = 10) -> RoomMessage:
        matches = provider.search(
            what=(prof.company or prof.query or (prof.titles or [""])[0]),
            where=prof.location, company=prof.company, limit=limit,
        )
        for m in matches:
            m.fit_score, m.fit_reasons = self._score(prof, m)
        matches.sort(key=lambda m: m.fit_score, reverse=True)
        jm = JobMatches(profile_query=prof.query, provider=provider.name,
                        count=len(matches), matches=matches)
        if matches:
            top = matches[0]
            line = (f"Found {len(matches)} live postings via {provider.name}. "
                    f"Best fit: {top.title} @ {top.company} ({top.fit_score:.0%}). "
                    "@commander, present these for approval.")
        else:
            line = f"No live postings matched '{prof.query}' on {provider.name}. @commander."
        return RoomMessage.of(self.id, Intent.JOB_MATCHES, self._voice("rank", line),
                              mentions=["@commander"], payload_model=jm)

    @staticmethod
    def _score(prof: SearchProfile, m: JobMatch) -> tuple[float, list[str]]:
        text = f"{m.title} {m.description}".lower()
        reasons: list[str] = []
        skills = [s.lower() for s in prof.skills]
        matched = [s for s in skills if re.search(r"\b" + re.escape(s) + r"\b", text)]
        score = 0.0
        if skills:
            score += 0.7 * (len(matched) / max(1, len(skills)))
            if matched:
                reasons.append(f"skills overlap: {', '.join(matched[:6])}")
        # Title / field alignment
        q_terms = re.findall(r"[a-z]+", (prof.query or "").lower())
        title_hits = [t for t in q_terms if len(t) > 2 and t in text]
        if title_hits:
            score += 0.3 * min(1.0, len(title_hits) / max(1, len(q_terms)))
            reasons.append(f"role match: {' '.join(title_hits[:4])}")
        if not skills and title_hits:
            score = max(score, 0.45 + 0.1 * len(title_hits))
        if not reasons:
            reasons.append("listed as a current opening for the query")
            score = max(score, 0.25)
        return round(min(1.0, score), 3), reasons


# --------------------------------------------------------------------------- #
class Commander(JobAgent):
    framework, provider = "orchestrator", "aiml"

    def request_approval(self, jm: JobMatches) -> RoomMessage:
        line = (f"{jm.count} ranked matches ready. Approve to tailor + apply to the "
                "selected role? @human [y/N]")
        return RoomMessage.of(self.id, Intent.APPROVAL_REQUEST, self._voice("ask", line),
                              mentions=["@human"])

    def decide(self, approved_by: str, proceeded: bool, chosen_id: str | None,
               submitted: int, queued: int) -> RoomMessage:
        dec = JobDecision(approved_by=approved_by, proceeded=proceeded,
                          chosen_match_id=chosen_id, submitted=submitted, queued=queued)
        if not proceeded:
            line = "Run halted by human — no applications prepared."
        else:
            line = (f"Approved by {approved_by}. {submitted} submitted, {queued} queued "
                    "with ready packages.")
        return RoomMessage.of(self.id, Intent.DECISION, self._voice("decide", line),
                              payload_model=dec)


# --------------------------------------------------------------------------- #
class Tailor(JobAgent):
    framework, provider = "LangGraph", "aiml"

    def tailor(self, prof: SearchProfile, match: JobMatch, do_it: bool, run_id: str) -> RoomMessage:
        if not do_it:
            tr = TailorResult(match_id=match.id, tailored=False,
                              note="User declined tailoring; original resume will be used.")
            return RoomMessage.of(self.id, Intent.TAILOR_RESULT,
                                  self._voice("skip", "Skipping resume tailoring as requested."),
                                  mentions=["@applier"], payload_model=tr)

        jd = f"{match.title} {match.description}".lower()
        posting_kw = sorted({s for s in SKILL_VOCAB if re.search(r"\b" + re.escape(s) + r"\b", jd)})
        have = {s.lower() for s in prof.skills}
        added = [k for k in posting_kw if k not in have]

        md = self._render_markdown(prof, match, posting_kw, added)
        base = ARTIFACTS / run_id
        files = _write_resume_files(base, f"resume_{match.id or 'role'}", md)

        tr = TailorResult(match_id=match.id, tailored=True, keywords_added=added,
                          markdown=md, files=files,
                          note=f"Aligned to '{match.title}' @ {match.company}; "
                               f"emphasized {len(posting_kw)} posting keywords.")
        line = (f"Tailored your resume to {match.title} @ {match.company} — "
                f"added {len(added)} keyword(s); md/PDF/DOCX ready. @applier.")
        return RoomMessage.of(self.id, Intent.TAILOR_RESULT, self._voice("tailor", line),
                              mentions=["@applier"], payload_model=tr)

    def _render_markdown(self, prof, match, posting_kw, added) -> str:
        skills = sorted(set([s.lower() for s in prof.skills] + posting_kw))
        summary = (prof.summary or "Experienced professional.").strip()
        tailored_summary = (
            f"{summary} Targeting **{match.title}** at **{match.company}**, with strength in "
            f"{', '.join(posting_kw[:6]) or 'the listed requirements'}."
        )
        live = "" if _offline() else self.llm.complete(
            system="Rewrite this resume summary to match a job posting. 2 sentences max.",
            user=f"Posting: {match.title} at {match.company}. {match.description[:600]}\n\n"
                 f"Candidate summary: {summary}", role="@tailor", tag="rewrite",
        )
        if live and live.strip() and live.strip() != summary:
            tailored_summary = live.strip()
        lines = [
            f"# {(prof.titles or [match.title])[0]}",
            "",
            f"_Tailored for {match.title} — {match.company}_",
            "",
            "## Summary",
            tailored_summary,
            "",
            "## Skills",
            ", ".join(skills) or "—",
            "",
            "## Experience",
            (prof.resume_text or "See attached.").strip()[:4000],
        ]
        return "\n".join(lines)


# --------------------------------------------------------------------------- #
class Applier(JobAgent):
    """Prepares the package and applies for real, or queues — never fakes 'applied'."""
    framework, provider = "CrewAI", "featherless"

    def apply(self, prof: SearchProfile, match: JobMatch, tailor: TailorResult,
              run_id: str) -> RoomMessage:
        cover = self._cover_letter(prof, match)
        resume_files = tailor.files if tailor.tailored else {}

        if match.apply_email:
            status, method, detail = self._email_apply(match, cover, resume_files)
        else:
            status, method = "queued", "apply_link"
            detail = ("No programmatic apply for this posting — queued with a one-click "
                      f"apply link and the ready package: {match.url}")

        app = Application(
            match_id=match.id, title=match.title, company=match.company,
            status=status, method=method, apply_url=match.url,
            apply_email=match.apply_email, cover_letter=cover,
            resume_files=resume_files, detail=detail,
        )
        verb = "Submitted" if status == "submitted" else "Queued"
        line = f"{verb}: {match.title} @ {match.company}. {detail[:80]}"
        return RoomMessage.of(self.id, Intent.APPLICATION, self._voice("apply", line),
                              payload_model=app)

    def _email_apply(self, match, cover, resume_files):
        from backend import email as mailer
        if not mailer.is_configured():
            return ("queued", "apply_link",
                    f"Email-apply address available ({match.apply_email}) but email is "
                    "not configured — queued instead of submitting (honest).")
        try:
            subject = f"Application: {match.title}"
            html = f"<p>{cover.replace(chr(10), '<br>')}</p>"
            mailer.send_email(subject, cover, html, to=[match.apply_email])
            return ("submitted", "email",
                    f"Application emailed to {match.apply_email} for {match.title}.")
        except Exception as e:
            return ("queued", "apply_link", f"Email-apply failed ({e}); queued with link.")

    def _cover_letter(self, prof: SearchProfile, match: JobMatch) -> str:
        top = ", ".join(prof.skills[:5]) or "relevant experience"
        fallback = (
            f"Dear {match.company} Hiring Team,\n\n"
            f"I'm excited to apply for the {match.title} role. My background in {top} "
            f"aligns closely with what you're looking for, and I'd welcome the chance to "
            f"contribute to your team.\n\nBest regards,\nThe candidate"
        )
        live = "" if _offline() else self.llm.complete(
            system="Write a concise 4-sentence cover letter. Professional, specific.",
            user=f"Role: {match.title} at {match.company}. Candidate skills: {top}. "
                 f"Posting: {match.description[:400]}", role="@applier", tag="cover",
        )
        return live.strip() if live and len(live.strip()) > 40 else fallback


# --------------------------------------------------------------------------- #
# Artifact rendering (md + PDF + DOCX)
# --------------------------------------------------------------------------- #
def _write_resume_files(base: Path, stem: str, markdown: str) -> dict[str, str]:
    base.mkdir(parents=True, exist_ok=True)
    files: dict[str, str] = {}

    md_path = base / f"{stem}.md"
    md_path.write_text(markdown, encoding="utf-8")
    files["md"] = str(md_path)

    try:
        files["pdf"] = _render_pdf(base / f"{stem}.pdf", markdown)
    except Exception:
        pass
    try:
        files["docx"] = _render_docx(base / f"{stem}.docx", markdown)
    except Exception:
        pass
    return files


def _render_pdf(path: Path, markdown: str) -> str:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    flow = []
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line:
            flow.append(Spacer(1, 6)); continue
        if line.startswith("# "):
            flow.append(Paragraph(line[2:], styles["Title"]))
        elif line.startswith("## "):
            flow.append(Paragraph(line[3:], styles["Heading2"]))
        else:
            safe = line.replace("**", "").replace("&", "&amp;").replace("<", "&lt;")
            flow.append(Paragraph(safe, styles["BodyText"]))
    doc.build(flow)
    return str(path)


def _render_docx(path: Path, markdown: str) -> str:
    import docx
    d = docx.Document()
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("# "):
            d.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            d.add_heading(line[3:], level=1)
        elif line:
            d.add_paragraph(line.replace("**", ""))
    d.save(str(path))
    return str(path)
